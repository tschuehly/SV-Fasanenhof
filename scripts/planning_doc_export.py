#!/usr/bin/env python3

from __future__ import annotations

import os
import sys
from pathlib import Path

BUNDLED_PYTHON = Path(
    "/Users/tschuehly/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3"
)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ModuleNotFoundError:
    if BUNDLED_PYTHON.exists() and Path(sys.executable) != BUNDLED_PYTHON:
        os.execv(str(BUNDLED_PYTHON), [str(BUNDLED_PYTHON), __file__, *sys.argv[1:]])
    raise


MARKDOWN_PATH = Path("/Users/tschuehly/IdeaProjects/SV-Fasanenhof/site-structure.md")
DOCX_PATH = Path(
    "/Users/tschuehly/Library/CloudStorage/OneDrive-Personal/Dokumente/Hobby/SV-Fasanenhof-Homepage.docx"
)

DOC_FONT = "Calibri"
GREEN_DARK = RGBColor(12, 74, 55)
GREEN = RGBColor(25, 104, 82)
GREEN_MUTED = RGBColor(48, 122, 100)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(89, 103, 99)


def set_font(style, name=DOC_FONT, size=None, bold=None, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def set_border(paragraph, bottom_color="7DB595", bottom_size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), bottom_size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), bottom_color)
    p_bdr.append(bottom)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    set_font(styles["Normal"], size=11.3, color=TEXT)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 28, GREEN_DARK, 0, 10),
        ("Heading 1", 22, GREEN_DARK, 18, 8),
        ("Heading 2", 17, GREEN, 13, 5),
        ("Heading 3", 13.2, GREEN_MUTED, 9, 4),
    ]:
        style = styles[name]
        set_font(style, size=size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if name.startswith("Heading"):
            style.paragraph_format.keep_with_next = True

    if "Meta" not in styles:
        meta = styles.add_style("Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["Meta"]
    set_font(meta, size=10, color=MUTED)
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.line_spacing = 1.0

    for list_style_name in ["List Bullet", "List Bullet 2", "List Bullet 3"]:
        if list_style_name in styles:
            style = styles[list_style_name]
            set_font(style, size=11.2, color=TEXT)
            style.paragraph_format.space_after = Pt(2.5)
            style.paragraph_format.line_spacing = 1.08


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.add_run("Webseitenstruktur und Inhaltsplan")
    p.paragraph_format.space_after = Pt(6)
    set_border(p)

    subtitle = doc.add_paragraph(
        "Arbeitsstruktur für die Erstellung und Erweiterung der SV-Fasanenhof-Webseite"
    )
    subtitle.runs[0].font.name = DOC_FONT
    subtitle.runs[0].font.size = Pt(12.5)
    subtitle.runs[0].font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)


def add_field_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = DOC_FONT
    r.font.size = Pt(11.4)
    r.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.left_indent = Cm(0.05)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.name = DOC_FONT
        p.runs[0].font.size = Pt(11.3)
        p.runs[0].font.color.rgb = TEXT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.left_indent = Cm(0.18)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.72)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.name = DOC_FONT
        run.font.size = Pt(11.2)
        run.font.color.rgb = TEXT


def export_markdown_to_docx(markdown_path: Path, docx_path: Path) -> None:
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown-Datei nicht gefunden: {markdown_path}")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    configure_styles(doc)
    add_title_block(doc)

    footer = section.footer.paragraphs[0]
    footer.text = "Webseitenstruktur und Inhaltsplan"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        footer.runs[0].font.name = DOC_FONT
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    seen_h1 = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph_buffer() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer if part.strip()))
            paragraph_buffer = []

    for raw in markdown_path.read_text(encoding="utf-8").splitlines():
        stripped = raw.strip()
        if not stripped:
            flush_paragraph_buffer()
            continue

        if stripped.startswith("# "):
            flush_paragraph_buffer()
            seen_h1 += 1
            heading = doc.add_heading(stripped[2:].strip(), level=1)
            if seen_h1 > 1:
                heading.paragraph_format.page_break_before = True
            set_border(heading, bottom_color="D9E7DD", bottom_size="6")
            continue

        if stripped.startswith("## "):
            flush_paragraph_buffer()
            doc.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("### "):
            flush_paragraph_buffer()
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("- "):
            flush_paragraph_buffer()
            add_bullet(doc, stripped[2:].strip())
            continue

        if stripped.endswith(":") and len(stripped) < 90:
            flush_paragraph_buffer()
            add_field_label(doc, stripped)
            continue

        add_line = stripped.replace("  ", " ")
        paragraph_buffer.append(add_line)

    flush_paragraph_buffer()

    props = doc.core_properties
    props.title = "Webseitenstruktur und Inhaltsplan"
    props.subject = "Strukturierte Auswertung für die Erstellung der SV-Fasanenhof-Webseite"
    props.author = "Codex"
    props.keywords = "SV Fasanenhof, Sitemap, Webseitenstruktur, Zielgruppen, Ziele"

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


if __name__ == "__main__":
    export_markdown_to_docx(MARKDOWN_PATH, DOCX_PATH)
